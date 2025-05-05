import random
from docx import Document

def shuffle_questions_in_docx(input_file, output_file):
    # Load the Word document
    doc = Document(input_file)
    
    # Collect all the questions in a list
    questions = []
    question = ""
    
    # Iterate through the paragraphs in the document and group the questions
    for para in doc.paragraphs:
        if para.text.strip() == "":
            continue  # Skip empty paragraphs
        
        # Detect question and answer block
        if para.text.startswith("Answer:") or para.text.startswith("•") or para.text.startswith("A)") or para.text.startswith("B)") or para.text.startswith("C)") or para.text.startswith("D)"):
            question += para.text.strip() + " "
        else:
            # If it's a new question or it's the first question, save the old one and reset
            if question:
                questions.append(question.strip())
            question = para.text.strip() + " "
    
    # Append the last question if it exists
    if question:
        questions.append(question.strip())

    # Shuffle the questions randomly
    random.shuffle(questions)

    # Create a new document to save the shuffled questions
    new_doc = Document()
    for index, q in enumerate(questions, start=1):
        new_doc.add_paragraph(f"{index}. {q}")  # Add sequential number before each question
    
    # Save the new shuffled document
    new_doc.save(output_file)

# Usage example
input_file = 'it all question.docx'  # Path to your input Word document
output_file = 'shuffle.docx'  # Path where the shuffled document will be saved

shuffle_questions_in_docx(input_file, output_file)

print(f"Questions shuffled and saved to {output_file}")
